"""
Convert VENA_API_PERFORMANCE_STRATEGY.md to Word document.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def convert_markdown_to_word():
    """Convert the performance strategy markdown to Word format."""
    
    # Read the markdown file
    with open('VENA_API_PERFORMANCE_STRATEGY.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines (we'll add spacing manually)
        if not line:
            i += 1
            continue
        
        # Handle headers
        if line.startswith('# '):
            # Main title
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            # Section heading
            heading = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Subsection heading
            heading = doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text (remove markdown)
            text = line.replace('**', '')
            para = doc.add_paragraph()
            run = para.add_run(text)
            run.bold = True
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            text = line[2:].strip()
            # Remove markdown formatting
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            para = doc.add_paragraph(text, style='List Bullet')
        elif any(line.startswith(f'{num}. ') for num in range(1, 10)):
            # Numbered list
            text = line.split('. ', 1)[1] if '. ' in line else line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            para = doc.add_paragraph(text, style='List Number')
        elif line.startswith('```'):
            # Code block - skip the opening/closing markers
            i += 1
            continue
        else:
            # Regular paragraph
            # Remove markdown formatting
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            text = re.sub(r'`(.*?)`', r'\1', text)
            para = doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    output_file = 'VENA_API_PERFORMANCE_STRATEGY.docx'
    doc.save(output_file)
    print(f"Successfully converted to {output_file}")

if __name__ == "__main__":
    convert_markdown_to_word()
