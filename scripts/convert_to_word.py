"""
Convert ARCHITECTURE.md to Word document format.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path
import re

def create_word_document(md_file_path: str, output_path: str):
    """Convert markdown file to Word document."""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read markdown file
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Main title (H1)
        if line.startswith('# ') and not line.startswith('##'):
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            i += 1
            continue
        
        # H2 headings
        if line.startswith('## ') and not line.startswith('###'):
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            i += 1
            continue
        
        # H3 headings
        if line.startswith('### '):
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            i += 1
            continue
        
        # H4 headings
        if line.startswith('#### '):
            title = line[5:].strip()
            p = doc.add_heading(title, level=4)
            i += 1
            continue
        
        # Horizontal rule
        if line.startswith('---'):
            p = doc.add_paragraph()
            p.add_run('─' * 50).bold = True
            i += 1
            continue
        
        # Table
        if '|' in line and '---' in lines[i+1] if i+1 < len(lines) else False:
            # Parse table
            table_lines = []
            j = i
            while j < len(lines) and '|' in lines[j]:
                table_lines.append(lines[j])
                j += 1
            
            if len(table_lines) >= 2:
                # Create table
                rows = len(table_lines) - 1  # Exclude separator row
                cols = len([c for c in table_lines[0].split('|') if c.strip()])
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                # Header row
                header_cells = [c.strip() for c in table_lines[0].split('|') if c.strip()]
                for idx, cell in enumerate(header_cells):
                    table.rows[0].cells[idx].text = cell
                    table.rows[0].cells[idx].paragraphs[0].runs[0].bold = True
                
                # Data rows
                for row_idx, table_line in enumerate(table_lines[2:], start=1):
                    cells = [c.strip() for c in table_line.split('|') if c.strip()]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < len(table.rows[row_idx].cells):
                            table.rows[row_idx].cells[col_idx].text = cell_text
                
                i = j
                continue
        
        # Bullet list item
        if line.startswith('- '):
            text = line[2:].strip()
            # Handle bold text
            p = doc.add_paragraph(text, style='List Bullet')
            format_bold_in_paragraph(p, text)
            i += 1
            continue
        
        # Numbered list item
        if re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line).strip()
            p = doc.add_paragraph(text, style='List Number')
            format_bold_in_paragraph(p, text)
            i += 1
            continue
        
        # Regular paragraph
        text = line
        # Handle bold markers
        p = doc.add_paragraph()
        format_text_with_markdown(p, text)
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"Word document created: {output_path}")

def format_bold_in_paragraph(paragraph, text):
    """Format bold text in paragraph."""
    # Remove existing runs
    paragraph.clear()
    
    # Find bold markers
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def format_text_with_markdown(paragraph, text):
    """Format text with markdown syntax."""
    # Handle code blocks (backticks)
    parts = re.split(r'(`[^`]+`)', text)
    for part in parts:
        if part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.highlight_color = 7  # Yellow highlight
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

if __name__ == '__main__':
    import sys
    project_root = Path(__file__).parent.parent
    
    # Accept command line arguments or use defaults
    if len(sys.argv) >= 3:
        md_file = project_root / sys.argv[1]
        output_file = project_root / sys.argv[2]
    else:
        md_file = project_root / 'ARCHITECTURE.md'
        output_file = project_root / 'ARCHITECTURE.docx'
    
    if not md_file.exists():
        print(f"Error: {md_file} not found!")
        exit(1)
    
    create_word_document(str(md_file), str(output_file))
    print(f"\nWord document saved to: {output_file}")
